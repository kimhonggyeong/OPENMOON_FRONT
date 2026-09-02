from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "OPENMOON_AI_사용방법.docx"
LOGO = ROOT / "frontend" / "public" / "yullinmoon-logo.png"

FONT = "맑은 고딕"
ORANGE = "D96D31"
DARK = "2F2925"
MUTED = "756C64"
PALE = "F7F2EC"
PALE_ORANGE = "FFF2E8"
BORDER = "DED7CE"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=130, bottom=100, end=130) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_no_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_run(run, size=11, bold=False, color=DARK, italic=False) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_text(doc, text, *, size=11, bold=False, color=DARK, after=6, before=0, align=None, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if align is not None:
        p.alignment = align
    set_run(p.add_run(text), size=size, bold=bold, color=color, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt({1: 18, 2: 14, 3: 10}[level])
    p.paragraph_format.space_after = Pt({1: 10, 2: 7, 3: 5}[level])
    set_run(p.add_run(text), size={1: 16, 2: 13, 3: 11.5}[level], bold=True, color=ORANGE if level < 3 else DARK)
    return p


def add_step(doc, number: str, title: str, body: str, note: str | None = None):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(0.72)
    table.columns[1].width = Inches(5.78)
    set_table_borders(table, color="E6DED5", size="5")
    left, right = table.rows[0].cells
    left.width = Inches(0.72)
    right.width = Inches(5.78)
    set_cell_shading(left, ORANGE)
    set_cell_shading(right, WHITE)
    for cell in (left, right):
        set_cell_margins(cell, 120, 140, 120, 140)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    lp = left.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(lp.add_run(number), size=12, bold=True, color=WHITE)
    rp = right.paragraphs[0]
    rp.paragraph_format.space_after = Pt(3)
    set_run(rp.add_run(title), size=11.5, bold=True)
    rp2 = right.add_paragraph()
    rp2.paragraph_format.space_after = Pt(0)
    rp2.paragraph_format.line_spacing = 1.22
    set_run(rp2.add_run(body), size=10.2, color=DARK)
    if note:
        rp3 = right.add_paragraph()
        rp3.paragraph_format.space_before = Pt(5)
        rp3.paragraph_format.space_after = Pt(0)
        set_run(rp3.add_run(f"참고  {note}"), size=9.2, bold=True, color=ORANGE)
    add_text(doc, "", size=1, after=3)


def add_callout(doc, label: str, text: str, caution=False):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_shading(cell, "FFF0EE" if caution else PALE_ORANGE)
    set_cell_margins(cell, 140, 170, 140, 170)
    set_table_borders(table, color="E8B8B3" if caution else "F0C8AA", size="6")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    set_run(p.add_run(label), size=10, bold=True, color="A13831" if caution else ORANGE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.25
    set_run(p2.add_run(text), size=10, color=DARK)
    add_text(doc, "", size=1, after=4)


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    for idx, (cell, header, width) in enumerate(zip(table.rows[0].cells, headers, widths)):
        cell.width = Inches(width)
        set_cell_shading(cell, ORANGE)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(header), size=9.5, bold=True, color=WHITE)
    set_repeat_table_header(table.rows[0])
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        set_repeat_no_split(table.rows[-1])
        for idx, (cell, value, width) in enumerate(zip(cells, row, widths)):
            cell.width = Inches(width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if ridx % 2:
                set_cell_shading(cell, PALE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            set_run(p.add_run(value), size=9.2, bold=(idx == 0), color=DARK)
    add_text(doc, "", size=1, after=5)
    return table


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(4)
    set_run(p.add_run("OPENMOON AI 견적 업무 보조  |  "), size=8.5, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for level in (1, 2, 3):
        style = doc.styles[f"Heading {level}"]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt({1: 16, 2: 13, 3: 11.5}[level])
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(ORANGE if level < 3 else DARK)


def build():
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)
    add_footer(section)

    # Cover
    add_text(doc, "사용자 매뉴얼", size=11, bold=True, color=ORANGE, after=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(20)
        p.add_run().add_picture(str(LOGO), width=Inches(2.6))
    add_text(doc, "OPENMOON AI 견적 업무 보조", size=25, bold=True, color=DARK, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(doc, "메일 분석부터 견적서 생성·검토·발송까지", size=13, color=MUTED, after=28, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_callout(doc, "이 문서의 목적", "처음 사용하는 직원이 서버에 접속하고, 메일을 분석하고, 견적서를 검토하여 안전하게 발송할 수 있도록 실제 화면의 버튼 이름과 업무 순서에 맞춰 설명합니다.")
    add_heading(doc, "가장 빠른 시작", 2)
    add_step(doc, "1", "프로그램 실행", "배포받은 폴더 안의 OPENMOON_AI_LAN.exe를 실행합니다.")
    add_step(doc, "2", "서버 연결", "서버 컴퓨터는 ‘서버 열기’, 다른 컴퓨터는 ‘서버 찾기’를 선택합니다.")
    add_step(doc, "3", "사용자 선택", "이름과 색상을 선택하고 ‘선택 완료’를 누른 뒤 ‘웹 열기’를 누릅니다.")
    add_step(doc, "4", "업무 시작", "메일 동기화 → AI 분석 → 견적 검토 → 견적서 생성 → 발송 전 검토 순서로 진행합니다.")
    add_text(doc, "문서 버전 1.0  |  2026년 8월", size=8.5, color=MUTED, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_page_break(doc)
    add_heading(doc, "1. 실행 전 알아두기", 1)
    add_callout(doc, "별도 설치 불필요", "본 프로그램은 설치형 프로그램이 아닙니다. 전달받은 폴더 전체를 원하는 위치에 복사한 뒤 OPENMOON_AI_LAN.exe를 실행하면 됩니다.")
    add_heading(doc, "폴더 사용 주의사항", 2)
    add_matrix(doc, ["구분", "올바른 사용 방법"], [
        ("폴더 이동", "OPENMOON_AI_LAN 폴더 전체를 함께 이동합니다."),
        ("내부 폴더", "_internal, backend, data 폴더의 이름을 바꾸거나 일부만 이동하지 않습니다."),
        ("데이터", "업무 데이터는 서버 컴퓨터 기준으로 관리됩니다. 게스트 PC에 별도 데이터를 복사하지 않습니다."),
        ("메일 설정", "메일 계정 및 발송 설정은 관리자가 준비한 값을 사용합니다."),
    ], [1.45, 5.05])
    add_callout(doc, "중요", "한 사무실에서는 한 대의 컴퓨터만 공유 서버를 열고, 나머지 컴퓨터는 모두 게스트로 접속하는 것을 권장합니다.", caution=True)
    add_heading(doc, "프로그램 종료", 2)
    add_text(doc, "웹 브라우저 창만 닫으면 화면만 종료됩니다. 서버 공유까지 끝내려면 실행 창에서 ‘공유 서버 종료’ 또는 ‘서버 연결 종료’를 누릅니다.")

    add_page_break(doc)
    add_heading(doc, "2. 서버장과 게스트 접속", 1)
    add_matrix(doc, ["역할", "선택", "설명"], [
        ("서버장", "서버 열기", "업무 데이터와 견적 파일을 제공하는 기준 컴퓨터입니다."),
        ("게스트", "서버 찾기", "사내 네트워크에서 서버장을 찾아 같은 데이터로 업무합니다."),
    ], [1.1, 1.35, 4.05])
    add_heading(doc, "서버 컴퓨터", 2)
    add_step(doc, "1", "서버 열기", "프로그램 시작 화면에서 ‘서버 열기’를 선택합니다.")
    add_step(doc, "2", "사용자 정보", "사용자 이름과 본인을 구분할 색상을 선택합니다.")
    add_step(doc, "3", "웹 열기", "‘웹 열기’를 눌러 업무 화면을 엽니다. 서버 실행 창은 닫지 않습니다.")
    add_heading(doc, "게스트 컴퓨터", 2)
    add_step(doc, "1", "서버 찾기", "같은 사내 네트워크에서 ‘서버 찾기’를 선택합니다.")
    add_step(doc, "2", "서버 선택", "검색된 서버의 컴퓨터 이름과 주소를 확인하고 연결합니다.")
    add_step(doc, "3", "사용자 정보", "자신의 이름과 색상을 선택한 뒤 웹 화면을 엽니다.")
    add_callout(doc, "연결되지 않을 때", "서버 컴퓨터에서 프로그램이 실행 중인지, 두 컴퓨터가 같은 공유기 또는 사내 네트워크에 연결되어 있는지 먼저 확인합니다.")

    add_page_break(doc)
    add_heading(doc, "3. 메일 동기화와 AI 분석", 1)
    add_heading(doc, "메일 동기화", 2)
    add_step(doc, "1", "메일 화면 열기", "왼쪽 메뉴에서 ‘이메일’을 선택합니다.")
    add_step(doc, "2", "메일 동기화", "오른쪽 위 ‘메일 동기화’를 누릅니다. 새 메일이 목록에 들어올 때까지 기다립니다.")
    add_step(doc, "3", "메일 선택", "목록에서 처리할 메일을 선택해 원문, 발신자, 요청 내용을 확인합니다.")
    add_heading(doc, "AI 분석", 2)
    add_step(doc, "1", "분석 실행", "선택한 메일의 AI 분석 영역에서 ‘분석’을 누릅니다. 여러 건은 ‘전체 분석’으로 처리할 수 있습니다.")
    add_step(doc, "2", "분석 결과 확인", "고객명, 품목, 규격, 수량, 재질, 요청사항과 분석 요약을 확인합니다.")
    add_step(doc, "3", "누락·충돌 검토", "‘검토 필요’ 항목은 원문과 첨부파일을 확인하여 값이 맞는지 직접 검토합니다.")
    add_callout(doc, "AI 사용 원칙", "AI 분석 결과는 업무 보조 자료입니다. 고객이 보낸 원문, 첨부파일, 수량, 규격과 최종 금액은 담당자가 반드시 확인해야 합니다.", caution=True)
    add_heading(doc, "공용 표시", 2)
    add_text(doc, "별표는 개인 메일 표시이고, 하트는 사내 공용 표시입니다. 하트는 누른 사용자의 색으로 표시되며 모든 접속자에게 즉시 동기화됩니다.")

    add_page_break(doc)
    add_heading(doc, "4. 견적 내용 검토와 수정", 1)
    add_heading(doc, "품목 확인", 2)
    add_step(doc, "1", "품목 목록 확인", "AI가 추출한 품목명, 규격, 수량, 단가와 공급금액을 확인합니다.")
    add_step(doc, "2", "값 수정", "잘못된 값은 해당 입력칸에서 수정하고 ‘수정 저장’을 누릅니다.")
    add_step(doc, "3", "품목 추가", "메일에 여러 품목이 있으나 누락되었다면 ‘품목 추가’로 새 항목을 만듭니다.")
    add_heading(doc, "과거 견적과 단가 근거", 2)
    add_matrix(doc, ["영역", "용도"], [
        ("동일 회사 견적", "현재 고객사에 과거 발송한 견적 품목과 가격을 확인합니다."),
        ("동일 고객 과거 견적", "같은 고객의 이전 거래 내역을 확인합니다."),
        ("현재 단가 후보", "단가표와 과거 자료에서 검색된 가격 근거를 확인합니다."),
        ("근거 파일", "표시된 Excel 근거를 열어 원본 시트와 셀을 확인합니다."),
    ], [2.0, 4.5])
    add_heading(doc, "견적 에이전트", 2)
    add_text(doc, "오른쪽 견적 에이전트에 자연어로 질문할 수 있습니다. 예: ‘이 회사의 과거 현수막 단가를 찾아줘’, ‘현재 수량을 3개로 바꿔줘’, ‘단가표 파일을 열어줘’. 변경 요청 후에는 화면의 최종 값을 다시 확인합니다.")
    add_callout(doc, "협업 동기화", "서버장 또는 게스트가 품목을 수정하거나 분석하면 다른 접속자 화면에도 실시간으로 반영됩니다. 같은 항목을 동시에 수정하지 않도록 상단 접속자 목록을 확인합니다.")

    add_page_break(doc)
    add_heading(doc, "5. 견적서 생성", 1)
    add_step(doc, "1", "최종 값 확인", "품목, 규격, 수량, 단가, 공급금액과 총액을 확인합니다.")
    add_step(doc, "2", "견적서 생성", "AI 분석 영역 아래의 ‘견적서 생성’을 누릅니다.")
    add_step(doc, "3", "저장 방식 선택", "기존 파일, 부서 기준, 담당자 기준 또는 별도 파일 중 업무에 맞는 저장 방식을 선택합니다.")
    add_step(doc, "4", "파일 생성 확인", "Excel 내부 견적서와 고객 발송용 PDF가 생성되었는지 견적서 목록에서 확인합니다.")
    add_callout(doc, "데이터 반영 기준", "견적서를 생성하거나 수정 저장한 것만으로는 과거 견적 DB에 추가되지 않습니다. ‘승인 및 답장’ 발송이 성공한 견적만 과거 견적 DB에 반영됩니다.")
    add_heading(doc, "견적서 수정 시", 2)
    add_text(doc, "같은 메일에서 견적서를 다시 생성하면 기존 초안이 최신 내용으로 갱신됩니다. 금액이나 품목이 바뀌면 고객용 PDF도 다시 생성한 뒤 발송해야 합니다.")
    add_heading(doc, "파일 열기", 2)
    add_matrix(doc, ["버튼", "열리는 파일"], [
        ("Excel", "내부 확인 및 보관용 견적서 파일"),
        ("PDF", "고객에게 실제 첨부되는 견적서 파일"),
    ], [1.35, 5.15])

    add_page_break(doc)
    add_heading(doc, "6. 발송 전 검토와 승인 답장", 1)
    add_step(doc, "1", "발송 전 검토", "견적서 목록에서 해당 견적의 ‘발송 전 검토’를 누릅니다.")
    add_step(doc, "2", "실제 수신자 확인", "패널 상단의 실제 발송 주소를 확인합니다. 테스트 모드에서는 테스트 주소와 실제 고객 주소가 따로 표시됩니다.")
    add_step(doc, "3", "담당 직원 선택", "발송 담당자를 선택합니다. 직원별 부서, 이름, 직급, 전화번호가 정식 메일 본문에 반영됩니다.")
    add_step(doc, "4", "제목·본문 수정", "메일 제목과 본문을 읽고 필요한 내용을 수정합니다. ‘임시저장’을 누르면 다른 접속자도 같은 내용을 확인할 수 있습니다.")
    add_step(doc, "5", "PDF와 금액 확인", "‘PDF 미리보기’로 첨부파일을 열고 품목 수와 최종 금액을 다시 확인합니다.")
    add_step(doc, "6", "승인 및 답장", "모든 내용이 맞으면 ‘승인 및 답장’을 누르고 최종 수신 주소 확인창에서 발송을 승인합니다.")
    add_callout(doc, "발송 실패 시", "바로 재시도하지 말고 고객 수신함과 보낸메일함을 먼저 확인합니다. 서버 응답 전에 연결이 끊긴 경우 이미 발송되었을 수 있습니다.", caution=True)
    add_heading(doc, "발송 완료 후", 2)
    add_text(doc, "성공하면 상태가 ‘SENT(발송 완료)’로 바뀌고 견적 이력 DB에 추가됩니다. 발송에 실패하면 DB에 추가되지 않으며 오류 내용이 견적 카드에 표시됩니다.")

    add_page_break(doc)
    add_heading(doc, "7. 여러 사용자와 함께 사용", 1)
    add_heading(doc, "접속자 표시", 2)
    add_text(doc, "화면 오른쪽 위에는 현재 접속 중인 서버장과 게스트의 이름과 색상이 모두 표시됩니다. 정상 종료 시 바로 사라지며 비정상 종료 시 약 45초 후 목록에서 제거됩니다.")
    add_heading(doc, "실시간으로 공유되는 작업", 2)
    add_matrix(doc, ["작업", "공유 방식"], [
        ("메일 분석", "분석 완료 직후 전체 화면에 반영"),
        ("품목·견적 수정", "저장 성공 직후 전체 화면에 반영"),
        ("하트", "누른 사용자의 색과 이름으로 즉시 반영"),
        ("견적서 생성·삭제", "목록과 상태가 즉시 갱신"),
        ("견적 에이전트 대화", "같은 메일의 대화와 변경 결과를 공유"),
        ("발송 전 검토 내용", "제목·본문 임시저장 시 다른 사용자에게 반영"),
    ], [2.05, 4.45])
    add_callout(doc, "협업 권장사항", "같은 견적을 두 명이 동시에 편집하면 마지막에 저장된 값이 반영될 수 있습니다. 상단 접속자 목록과 공용 하트를 이용해 담당 중인 메일을 표시하는 것을 권장합니다.")
    add_heading(doc, "게스트의 파일 열기", 2)
    add_text(doc, "게스트가 과거 견적 Excel을 열면 서버 컴퓨터의 원본 파일을 내려받아 게스트 PC의 임시 폴더에서 엽니다. 원본 데이터의 기준은 항상 서버 컴퓨터입니다.")

    add_page_break(doc)
    add_heading(doc, "8. 설정과 데이터 관리", 1)
    add_heading(doc, "견적 이력 수동 업데이트", 2)
    add_text(doc, "설정 화면의 ‘견적 이력 업데이트’는 실제로 발송 완료된 견적만 이력 DB에 추가하거나 갱신합니다. 초안, 미승인, 발송 실패 견적은 포함하지 않습니다.")
    add_heading(doc, "단가표 관리", 2)
    add_text(doc, "단가표 관리에서는 품목명, 규격, 재질, 단위, 단가와 비고를 검색·추가·수정·삭제할 수 있습니다. 변경 사항은 다음 단가 검색부터 적용됩니다.")
    add_callout(doc, "관리자 작업", "메일 계정, API 키, DB 경로, 발송 안전장치와 data 폴더 파일은 일반 사용자가 변경하지 않습니다. 변경이 필요하면 프로그램 관리자에게 요청합니다.", caution=True)
    add_heading(doc, "테스트 발송 표시", 2)
    add_text(doc, "승인 테스트 모드에서는 실제 고객이 아니라 관리자가 지정한 테스트 주소로만 발송됩니다. 발송 전 검토 패널의 주황색 안내와 ‘실제 발송 주소’를 반드시 확인합니다.")

    add_page_break(doc)
    add_heading(doc, "9. 문제 해결", 1)
    add_matrix(doc, ["증상", "확인 방법"], [
        ("서버가 검색되지 않음", "서버 PC 실행 여부, 같은 네트워크 연결, 방화벽 허용 여부를 확인합니다."),
        ("웹 화면이 열리지 않음", "실행 창에서 서버 연결 상태를 확인하고 ‘웹 열기’를 다시 누릅니다."),
        ("메일이 보이지 않음", "‘메일 동기화’를 누르고 메일 계정 연결 상태를 확인합니다."),
        ("견적서 생성이 안 됨", "필수 품목, 수량, 단가와 검토 필요 항목을 확인합니다."),
        ("PDF가 없다고 표시됨", "견적서를 다시 생성하여 최신 고객용 PDF를 만듭니다."),
        ("발송이 차단됨", "테스트 모드 또는 실제 발송 허용 설정을 관리자에게 확인합니다."),
        ("발송 실패", "수신함·보낸메일함을 먼저 확인한 뒤 오류 문구를 관리자에게 전달합니다."),
        ("다른 사용자와 값이 다름", "서버 연결 상태를 확인하고 화면을 새로 열어 최신 상태를 다시 불러옵니다."),
    ], [2.05, 4.45])
    add_heading(doc, "관리자에게 전달할 정보", 2)
    add_text(doc, "문제가 계속되면 발생 시각, 서버장/게스트 여부, 선택한 메일 제목, 견적 번호, 화면에 나온 오류 문구와 화면 캡처를 함께 전달합니다.")
    add_callout(doc, "최종 확인", "고객에게 발송하기 전에는 수신자, 제목, 본문, 담당 직원, 첨부 PDF, 품목, 수량, 단가와 최종 금액을 모두 확인합니다.")
    add_text(doc, "OPENMOON AI 견적 업무 보조 사용자 매뉴얼 끝", size=9, bold=True, color=MUTED, before=20, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Metadata and save.
    props = doc.core_properties
    props.title = "OPENMOON AI 견적 업무 보조 사용자 매뉴얼"
    props.subject = "서버 접속, 메일 분석, 견적 생성, 발송 전 검토 및 협업 사용법"
    props.author = "열린문디자인"
    props.keywords = "OPENMOON, 견적, 사용자 매뉴얼"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
