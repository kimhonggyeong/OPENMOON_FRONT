from pathlib import Path

from docx import Document


path = Path(__file__).resolve().parents[1] / "OPENMOON_AI_사용방법.docx"
doc = Document(path)
all_text = "\n".join(p.text for p in doc.paragraphs)
all_text += "\n" + "\n".join(
    cell.text for table in doc.tables for row in table.rows for cell in row.cells
)
required = [
    "서버 열기",
    "서버 찾기",
    "메일 동기화",
    "AI 분석",
    "견적서 생성",
    "발송 전 검토",
    "승인 및 답장",
    "견적 이력 업데이트",
    "단가표 관리",
    "문제 해결",
]
section = doc.sections[0]

print(f"FILE_BYTES={path.stat().st_size}")
print(f"SECTIONS={len(doc.sections)}")
print(f"PARAGRAPHS={len(doc.paragraphs)}")
print(f"TABLES={len(doc.tables)}")
print(f"IMAGES={len(doc.inline_shapes)}")
print(f"HEADINGS={sum(1 for p in doc.paragraphs if p.style.name.startswith('Heading'))}")
print(f"PAGE_BREAKS={sum(p._p.xml.count('w:type=\"page\"') for p in doc.paragraphs)}")
print(f"MISSING={','.join(value for value in required if value not in all_text)}")
print(f"EMPTY_TABLE_CELLS={sum(1 for table in doc.tables for row in table.rows for cell in row.cells if not cell.text.strip())}")
print(f"PAGE_SIZE={section.page_width.inches:.2f}x{section.page_height.inches:.2f}")
print(
    "MARGINS="
    f"{section.top_margin.inches:.2f},"
    f"{section.right_margin.inches:.2f},"
    f"{section.bottom_margin.inches:.2f},"
    f"{section.left_margin.inches:.2f}"
)
